from __future__ import annotations
from typing import Dict, List, Optional
import csv




def export_per_frame_csv(
   path: str,
   avg_pressure_per_frame: List[float],
   estimated_vgrf_per_frame: List[float],
) -> Optional[str]:
   try:
       with open(path, "w", newline="", encoding="utf-8") as f:
           w = csv.writer(f)
           w.writerow(["frame", "avg_pressure", "estimated_vGRF"])
           n = max(len(avg_pressure_per_frame), len(estimated_vgrf_per_frame))
           for i in range(n):
               ap = avg_pressure_per_frame[i] if i < len(avg_pressure_per_frame) else ""
               vg = estimated_vgrf_per_frame[i] if i < len(estimated_vgrf_per_frame) else ""
               w.writerow(
                   [
                       i,
                       f"{ap:.6f}" if ap != "" else "",
                       f"{vg:.6f}" if vg != "" else "",
                   ]
               )
       return None
   except Exception as e:
       return f"Failed to write per-frame CSV: {e}"




def export_summary_docx(summary: Dict, path: str) -> Optional[str]:
    try:
        from docx import Document
    except Exception:
        return "python-docx not installed. Add `python-docx` to requirements."
    try:
        doc = Document()
        doc.add_heading("Sole Solutions Report", 0)
        doc.add_paragraph("Session Summary")

        def add_line(k: str) -> None:
            if k in summary:
                doc.add_paragraph(f"{k.replace('_', ' ').title()}: {summary[k]}")

        for k in (
            "frames",
            "sensors",
            "global_min",
            "global_max",
            "contact_time_frames",
            "contact_threshold",
            "pti",
            "dt",
        ):
            add_line(k)

        # Small per-frame preview
        doc.add_paragraph("")
        doc.add_paragraph("Per-Frame (first 20)")
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Frame"
        hdr[1].text = "Avg Pressure"
        hdr[2].text = "Estimated vGRF"

        ap = summary.get("avg_pressure_per_frame", [])[:20]
        vg = summary.get("estimated_vgrf_per_frame", [])[:20]
        for i in range(max(len(ap), len(vg))):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = f"{ap[i]:.3f}" if i < len(ap) else ""
            row[2].text = f"{vg[i]:.3f}" if i < len(vg) else ""

        # Optional: saved segments / frame selections
        segments = summary.get("segments") or []
        if segments:
            doc.add_paragraph("")
            doc.add_paragraph("Saved Segments")

            seg_table = doc.add_table(rows=1, cols=7)
            sh = seg_table.rows[0].cells
            sh[0].text = "Name"
            sh[1].text = "Start"
            sh[2].text = "End"
            sh[3].text = "Duration (s)"
            sh[4].text = "Peak Pressure (kPa)"
            sh[5].text = "Peak vGRF (N)"
            sh[6].text = "Impulse (N·s)"

            for seg in segments:
                row = seg_table.add_row().cells
                row[0].text = str(seg.get("name", ""))
                row[1].text = str(seg.get("start_frame", ""))
                row[2].text = str(seg.get("end_frame", ""))
                dur = seg.get("duration_s", 0.0)
                row[3].text = f"{dur:.3f}" if isinstance(dur, (int, float)) else ""
                pp = seg.get("peak_pressure_kpa", 0.0)
                row[4].text = f"{pp:.1f}" if isinstance(pp, (int, float)) else ""
                pv = seg.get("peak_vgrf_N", 0.0)
                row[5].text = f"{pv:.1f}" if isinstance(pv, (int, float)) else ""
                im = seg.get("impulse_Ns", 0.0)
                row[6].text = f"{im:.2f}" if isinstance(im, (int, float)) else ""

        doc.save(path)
        return None
    except Exception as e:
        return f"Failed to write DOCX: {e}"


def export_summary_pdf(summary: Dict, path: str) -> Optional[str]:
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch
    except Exception:
        return "reportlab not installed. Add `reportlab` to requirements."
    try:
        c = canvas.Canvas(path, pagesize=LETTER)
        width, height = LETTER
        y = height - 1 * inch

        c.setFont("Helvetica-Bold", 16)
        c.drawString(1 * inch, y, "Sole Solutions Report")
        y -= 0.4 * inch

        c.setFont("Helvetica", 11)
        c.drawString(1 * inch, y, "Session Summary")
        y -= 0.25 * inch

        def line(txt: str) -> None:
            nonlocal y
            c.drawString(1 * inch, y, txt)
            y -= 0.2 * inch
            if y < 1 * inch:
                c.showPage()
                y = height - 1 * inch
                c.setFont("Helvetica", 11)

        for k in (
            "frames",
            "sensors",
            "global_min",
            "global_max",
            "contact_time_frames",
            "contact_threshold",
            "pti",
            "dt",
        ):
            if k in summary:
                line(f"{k.replace('_', ' ').title()}: {summary[k]}")

        # Segments section
        segments = summary.get("segments") or []
        if segments:
            line("")
            line("Saved Segments:")
            for seg in segments:
                name = seg.get("name", "")
                sf = seg.get("start_frame", "")
                ef = seg.get("end_frame", "")
                dur = seg.get("duration_s", 0.0)
                pp = seg.get("peak_pressure_kpa", 0.0)
                pv = seg.get("peak_vgrf_N", 0.0)
                im = seg.get("impulse_Ns", 0.0)
                line(
                    f"- {name} [{sf}-{ef}], "
                    f"dur={dur:.3f}s, "
                    f"peakP={pp:.1f} kPa, "
                    f"peakF={pv:.1f} N, "
                    f"impulse={im:.2f} N·s"
                )

        c.showPage()
        c.save()
        return None
    except Exception as e:
        return f"Failed to write PDF: {e}"
